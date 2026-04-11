"""
File Forge — Wave 4
====================
File preprocessing and chunking layer for the FOAI pipeline.
Adapted from Intelligent-Internet/II-Commons text chunking patterns.

Supports: PDF, DOCX, TXT, MD, CSV
Does NOT import II-Commons' heavy ML dependencies — uses lightweight libs only.

API:
    preprocess(file_path) -> Document
    chunk(document, chunk_size) -> list[Chunk]
"""

from __future__ import annotations

import csv
import io
import os
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


@dataclass
class Document:
    """Structured representation of a preprocessed file."""

    file_path: str
    file_name: str
    file_type: str
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    size_bytes: int = 0
    extracted_at: str = ""

    def to_dict(self) -> dict[str, Any]:
        return {
            "file_path": self.file_path,
            "file_name": self.file_name,
            "file_type": self.file_type,
            "text": self.text,
            "metadata": self.metadata,
            "size_bytes": self.size_bytes,
            "extracted_at": self.extracted_at,
        }


@dataclass
class Chunk:
    """A chunk of text from a document."""

    text: str
    index: int
    start_char: int
    end_char: int
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "text": self.text,
            "index": self.index,
            "start_char": self.start_char,
            "end_char": self.end_char,
            "metadata": self.metadata,
        }


# ---------------------------------------------------------------------------
# Text extractors (one per supported file type)
# ---------------------------------------------------------------------------


def _extract_txt(file_path: str) -> tuple[str, dict[str, Any]]:
    """Extract text from .txt or .md files."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return text, {"encoding": "utf-8"}


def _extract_csv(file_path: str) -> tuple[str, dict[str, Any]]:
    """Extract text from .csv files — converts to readable rows."""
    rows: list[str] = []
    row_count = 0
    col_names: list[str] = []
    with open(file_path, "r", encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            if i == 0:
                col_names = row
            rows.append(" | ".join(row))
            row_count += 1
    text = "\n".join(rows)
    return text, {"row_count": row_count, "columns": col_names}


def _extract_pdf(file_path: str) -> tuple[str, dict[str, Any]]:
    """Extract text from PDF using PyMuPDF (fitz) if available, else fallback."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "PyMuPDF (fitz) is required for PDF support. Install: pip install pymupdf"
        )

    doc = fitz.open(file_path)
    pages: list[str] = []
    for page in doc:
        pages.append(page.get_text())
    text = "\n\n".join(pages)
    metadata = {"page_count": len(doc), "title": doc.metadata.get("title", "")}
    doc.close()
    return text, metadata


def _extract_docx(file_path: str) -> tuple[str, dict[str, Any]]:
    """Extract text from DOCX using python-docx if available."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX support. Install: pip install python-docx"
        )

    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    metadata = {
        "paragraph_count": len(paragraphs),
        "core_properties": {},
    }
    if doc.core_properties:
        cp = doc.core_properties
        metadata["core_properties"] = {
            "title": cp.title or "",
            "author": cp.author or "",
            "subject": cp.subject or "",
        }
    return text, metadata


# Dispatcher
_EXTRACTORS: dict[str, Any] = {
    ".txt": _extract_txt,
    ".md": _extract_txt,
    ".csv": _extract_csv,
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
}

SUPPORTED_EXTENSIONS = set(_EXTRACTORS.keys())


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def preprocess(file_path: str) -> Document:
    """
    Read a file, extract its text and metadata, return a structured Document.

    Raises:
        FileNotFoundError: if the file doesn't exist
        ValueError: if the file type is not supported
        ImportError: if an optional dependency is missing (e.g., PyMuPDF for PDF)
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    if ext not in _EXTRACTORS:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    extractor = _EXTRACTORS[ext]
    text, metadata = extractor(file_path)

    return Document(
        file_path=str(path.resolve()),
        file_name=path.name,
        file_type=ext,
        text=text,
        metadata=metadata,
        size_bytes=path.stat().st_size,
        extracted_at=datetime.now(timezone.utc).isoformat(),
    )


def chunk(
    document: Document,
    chunk_size: int = 2048,
    overlap: int = 200,
) -> list[Chunk]:
    """
    Split a Document's text into overlapping chunks suitable for embedding/search.

    Algorithm adapted from II-Commons lib/text.py sentence-aware chunking,
    simplified to use regex sentence splitting (no NLTK dependency).

    Args:
        document: The preprocessed Document
        chunk_size: Target size of each chunk in characters
        overlap: Number of characters to overlap between chunks

    Returns:
        List of Chunk objects with positional metadata
    """
    text = document.text
    if not text or not text.strip():
        return []

    # Sentence-aware splitting (regex, no NLTK needed)
    sentences = _split_sentences(text)

    chunks: list[Chunk] = []
    current_text = ""
    current_start = 0
    char_pos = 0

    for sentence in sentences:
        sentence_len = len(sentence)

        if len(current_text) + sentence_len > chunk_size and current_text:
            # Emit current chunk
            chunks.append(
                Chunk(
                    text=current_text.strip(),
                    index=len(chunks),
                    start_char=current_start,
                    end_char=current_start + len(current_text),
                    metadata={
                        "file_name": document.file_name,
                        "file_type": document.file_type,
                    },
                )
            )
            # Overlap: keep tail of current chunk
            if overlap > 0 and len(current_text) > overlap:
                overlap_text = current_text[-overlap:]
                current_start = current_start + len(current_text) - overlap
                current_text = overlap_text
            else:
                current_start = char_pos
                current_text = ""

        current_text += sentence
        char_pos += sentence_len

    # Final chunk
    if current_text.strip():
        chunks.append(
            Chunk(
                text=current_text.strip(),
                index=len(chunks),
                start_char=current_start,
                end_char=current_start + len(current_text),
                metadata={
                    "file_name": document.file_name,
                    "file_type": document.file_type,
                },
            )
        )

    return chunks


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

# Regex for sentence boundaries — handles common abbreviations
_SENTENCE_RE = re.compile(r"(?<=[.!?])\s+(?=[A-Z])")


def _split_sentences(text: str) -> list[str]:
    """Split text into sentences using regex. Preserves whitespace in output."""
    parts = _SENTENCE_RE.split(text)
    if not parts:
        return [text]

    # Re-attach the whitespace separators so char positions stay correct
    result: list[str] = []
    pos = 0
    for part in parts:
        idx = text.find(part, pos)
        if idx > pos:
            # Prepend gap (whitespace between sentences) to this part
            result.append(text[pos:idx] + part)
        else:
            result.append(part)
        pos = idx + len(part)

    # Catch any trailing text
    if pos < len(text):
        if result:
            result[-1] += text[pos:]
        else:
            result.append(text[pos:])

    return result
